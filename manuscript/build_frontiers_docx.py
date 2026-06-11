from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "frontiers_neuroinformatics_spiketrainpatterndetector.md"
TEMPLATE_PATH = ROOT / "frontiers_template" / "Frontiers_Word_Templates" / "Frontiers_Template.docx"
OUT_PATH = ROOT / "frontiers_neuroinformatics_spiketrainpatterndetector_frontiers.docx"


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_run_font(run, name: str = "Times New Roman", size: float = 12, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_core_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title.font.size = Pt(16)
    title.font.bold = True
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for name, size, after in [
        ("Heading 1", 14, 8),
        ("Heading 2", 12, 6),
        ("Heading 3", 12, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10 if name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if "Author List" in doc.styles:
        author = doc.styles["Author List"]
        author.font.name = "Times New Roman"
        author._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        author.font.size = Pt(12)
        author.paragraph_format.space_after = Pt(6)
        author.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(12)
    caption.font.bold = True
    caption.paragraph_format.space_before = Pt(8)
    caption.paragraph_format.space_after = Pt(4)


def add_line_numbering(section) -> None:
    sect_pr = section._sectPr
    for existing in list(sect_pr.findall(qn("w:lnNumType"))):
        sect_pr.remove(existing)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "newPage")
    sect_pr.append(ln)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, size=10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_sep)
    r._r.append(txt)
    r._r.append(fld_end)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        add_line_numbering(section)
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        add_page_number(p)


def add_inline_markdown(paragraph, text: str, base_size: float = 12) -> None:
    # Handles **bold** and `code` spans. This is enough for the manuscript draft.
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=base_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Courier New", size=base_size - 1)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size)


def add_para(doc: Document, text: str = "", style: str = "Normal", base_size: float = 12):
    p = doc.add_paragraph(style=style)
    add_inline_markdown(p, text, base_size=base_size)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def apply_numbering(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), str(level))
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2:
        header = rows[0]
        body = rows[2:] if all(re.match(r"^-+$|^:?-+:?$", c.replace(" ", "")) for c in rows[1]) else rows[1:]
        return header, body, i
    return [], [], i


def set_cell_text(cell, text: str, bold: bool = False, base_size: float = 12) -> None:
    text = text.replace("inst/extdata/Grechishnikova_STN_2017_subset.csv", "Bundled example CSV (inst/extdata)")
    cell.text = ""
    p = cell.paragraphs[0]
    add_inline_markdown(p, text, base_size=base_size)
    for run in p.runs:
        run.bold = bold or run.bold
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def set_table_cell_margins(table, top=80, bottom=80, start=120, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_word_table(doc: Document, caption: str, header: list[str], body: list[list[str]]) -> None:
    add_para(doc, caption, style="Caption")
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.autofit = False
    set_table_cell_margins(table)
    ncols = len(header)
    if ncols == 2:
        widths = [3.8, 2.7]
        cell_size = 12
    elif ncols == 4:
        widths = [1.35, 1.9, 1.55, 1.7]
        cell_size = 10
    elif ncols == 5:
        widths = [1.05, 1.75, 0.9, 1.25, 1.55]
        cell_size = 9
    else:
        widths = [6.5 / max(ncols, 1)] * ncols
        cell_size = 9
    for j, cell in enumerate(table.rows[0].cells):
        set_cell_text(cell, header[j], bold=True, base_size=cell_size)
    for row in body:
        cells = table.add_row().cells
        for j, cell in enumerate(cells):
            set_cell_text(cell, row[j] if j < len(row) else "", base_size=cell_size)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] if idx < len(widths) else widths[-1])
    doc.add_paragraph()


def clean_heading(text: str) -> str:
    heading = text.strip().lstrip("#").strip()
    return re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", heading)


def extract_first_matching(lines: list[str], label: str) -> str:
    for line in lines:
        if line.startswith(label):
            return line.split(":", 1)[1].strip()
    return ""


def clean_metadata_value(value: str) -> str:
    value = value.strip()
    # Remove only wrapper markup from the field label/value, not meaningful
    # content such as the corresponding-author asterisk in "Houchun Zhou1*".
    if value.startswith("**"):
        value = value[2:].lstrip()
    if value.endswith("**"):
        value = value[:-2].rstrip()
    return value.strip()


def count_words(md: str) -> int:
    text = re.sub(r"`[^`]+`", " ", md)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)
    return len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", text))


def build() -> None:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Frontiers template not found: {TEMPLATE_PATH}")
    md = MD_PATH.read_text(encoding="utf-8")
    lines = md.splitlines()

    title = clean_heading(next(line for line in lines if line.startswith("# ")))
    article_type = clean_metadata_value(extract_first_matching(lines, "**Article type:**"))
    target_journal = clean_metadata_value(extract_first_matching(lines, "**Target journal:**"))
    running_title = clean_metadata_value(extract_first_matching(lines, "**Running title:**"))
    authors = clean_metadata_value(extract_first_matching(lines, "**Authors:**"))
    keywords = clean_metadata_value(extract_first_matching(lines, "**Keywords:**"))
    figure_count = sum(1 for line in lines if line.startswith("**Figure "))
    table_count = sum(1 for line in lines if line.startswith("**Table "))
    word_count = count_words(md)

    doc = Document(TEMPLATE_PATH)
    clear_document_body(doc)
    set_core_styles(doc)
    configure_sections(doc)

    add_para(doc, title, style="Title")
    add_para(doc, authors, style="Author List")
    add_para(doc, "1. [Institution, Department, City, Country - to be completed before submission]")
    add_para(doc, "* Correspondence: Houchun Zhou, zhouhouchun@outlook.com")
    add_para(doc, f"Article type: {article_type}")
    add_para(doc, f"Target journal: {target_journal}")
    add_para(doc, f"Running title: {running_title}")
    add_para(doc, f"Word count: {word_count}; Figures: {figure_count}; Tables: {table_count}")
    add_para(doc, f"Keywords: {keywords}", style="Author List")

    abstract_start = lines.index("## Abstract") + 1
    intro_start = next(i for i, line in enumerate(lines) if line.startswith("## 1. "))
    add_para(doc, "Abstract", style="Author List")
    for line in lines[abstract_start:intro_start]:
        if line.strip():
            add_para(doc, line.strip())

    tables_to_append: list[tuple[str, list[str], list[list[str]]]] = []
    i = intro_start
    in_references = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("**Table "):
            caption = re.sub(r"^\*\*|\*\*$", "", stripped)
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and lines[j].strip().startswith("|"):
                header, body, next_i = parse_markdown_table(lines, j)
                tables_to_append.append((caption, header, body))
                add_para(doc, f"{caption} is provided in the Tables section at the end of the manuscript.")
                i = next_i
                continue

        if stripped.startswith("## "):
            heading = clean_heading(stripped)
            in_references = heading == "References"
            add_para(doc, heading, style="Heading 1")
        elif stripped.startswith("### "):
            add_para(doc, clean_heading(stripped), style="Heading 2")
        elif stripped.startswith("#### "):
            add_para(doc, clean_heading(stripped), style="Heading 3")
        elif re.match(r"^\d+\.\s+", stripped):
            p = add_para(doc, re.sub(r"^\d+\.\s+", "", stripped), style="List Paragraph")
            apply_numbering(p, num_id=15, level=0)
            p.paragraph_format.space_after = Pt(4)
        elif stripped.startswith("- "):
            p = add_para(doc, stripped[2:], style="List Paragraph")
            apply_numbering(p, num_id=3, level=0)
            p.paragraph_format.space_after = Pt(4)
        elif stripped.startswith("|"):
            header, body, next_i = parse_markdown_table(lines, i)
            add_word_table(doc, "Table", header, body)
            i = next_i
            continue
        else:
            p = add_para(doc, stripped)
            if in_references and re.match(r"^[A-Z].*\([0-9]{4}\)\.", stripped):
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.paragraph_format.space_after = Pt(6)
        i += 1

    if tables_to_append:
        doc.add_page_break()
        add_para(doc, "Tables", style="Heading 1")
        for caption, header, body in tables_to_append:
            add_word_table(doc, caption, header, body)

    doc.core_properties.title = title
    doc.core_properties.author = "Houchun Zhou"
    doc.core_properties.subject = "Frontiers in Neuroinformatics Technology and Code manuscript"
    doc.core_properties.keywords = keywords
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
