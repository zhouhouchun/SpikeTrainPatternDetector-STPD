from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
BASE_MD_PATH = ROOT / "frontiers_neuroinformatics_spiketrainpatterndetector.md"
OUT_MD_PATH = ROOT / "neuroinformatics_software_original_article_rewritten.md"
OUT_PATH = Path("/Users/zark/Desktop/spiketrainpatterndetector_neuroinformatics_software_original_article_rewritten.docx")

BASE_FONT = "Times New Roman"
BASE_SIZE = 10


def transform_markdown(md: str) -> str:
    """Retarget and strengthen the existing draft for Neuroinformatics."""
    md = md.replace(
        "# SpikeTrainPatternDetector: an auditable R/Shiny framework for spike-train event review and validation",
        "# SpikeTrainPatternDetector: an auditable R/Shiny software framework for spike-train event review and validation",
    )
    md = md.replace("**Article type:** Technology and Code", "**Article type:** Software Original Article")
    md = md.replace("**Target journal:** Frontiers in Neuroinformatics", "**Target journal:** Neuroinformatics")
    md = md.replace(
        "**Keywords:** spike train, burst detection, pause detection, event grammar, neural manifold, dimensionality reduction, Shiny, reproducible neuroinformatics",
        "**Keywords:** spike train, burst detection, pause detection, event grammar, R/Shiny, neuroinformatics",
    )
    md = md.replace(
        "Here we describe the software architecture, algorithmic logic, validation outputs, optional exploratory modules, and example use of SpikeTrainPatternDetector. We frame the package as a Technology and Code contribution for neuroinformatics: a reproducible, human-readable, open software implementation designed to support spike-train analysis rather than to replace expert scientific judgment.",
        "Here we describe the software architecture, algorithmic logic, validation outputs, optional exploratory modules, and example use of SpikeTrainPatternDetector. We frame the package as a Software Original Article for Neuroinformatics: a reproducible, human-readable, open software implementation designed to support spike-train analysis without replacing expert scientific judgment.",
    )
    md = md.replace(
        "For final Frontiers submission, the project should be deposited in a stable public repository with a persistent DOI or URI.",
        "For final Neuroinformatics submission, the project should be deposited in a stable public repository with a persistent DOI or URI.",
    )
    md = md.replace(
        "The grammar uses the following conceptual layers.",
        "The grammar uses the following conceptual layers. Formally, for a train with ordered timestamps t_1, ..., t_n, the software recomputes timestamp-derived intervals Delta_i = t_i - t_{i-1}. A candidate c is an adjacent ISI interval [a, b], not an isolated point label. For each candidate, the detector constructs an evidence vector E(c) containing seed compactness, bridge burden, flanking gaps, boundary contrast, local variation, refractory or artifact flags, manual-lock status, and state-family features. The final public label is a deterministic function of E(c), the frozen threshold table, and the manual-arbitration policy. This representation is central to the design: the algorithm first exposes evidence about a candidate interval, then assigns a reviewable state label, and finally exports enough information for the decision to be audited or challenged.",
    )
    md = md.replace(
        "according to Frontiers and ICMJE authorship criteria",
        "according to ICMJE and Neuroinformatics authorship criteria",
    )
    md = md.replace(
        "according to Frontiers policy at the time of submission",
        "according to Springer Nature and Neuroinformatics policy at the time of submission",
    )

    md = replace_sections(
        md,
        skip_headings={
            "Data availability statement",
            "Code availability statement",
            "Ethics statement",
            "Author contributions",
            "Funding",
            "Conflict of interest",
            "Generative AI statement",
        },
    )
    md = md.replace("## Acknowledgments", declarations_block() + "\n\n## Acknowledgments")
    md = md.replace("## Figure legends", "## Figure captions")
    md = md.replace("**Figure 1.", "**Fig. 1.")
    md = md.replace("**Figure 2.", "**Fig. 2.")
    md = md.replace("**Figure 3.", "**Fig. 3.")
    md = md.replace("**Figure 4.", "**Fig. 4.")
    md = md.replace("**Figure 5.", "**Fig. 5.")
    md = re.sub(r"\*\*Fig\. ([0-9]+)\. ([^*]+)\.\*\*", r"**Fig. \1** \2.", md)
    md = re.sub(r"doi: ([0-9][^ \n]+)", r"https://doi.org/\1", md)
    return md


def declarations_block() -> str:
    return """## Statements and Declarations

### Ethics approval and consent to participate

This manuscript draft describes software and a bundled smoke-test dataset. It does not report new human or animal experimental results. If the software is submitted together with patient recordings, intraoperative data, animal experiments, or unpublished laboratory datasets, the final manuscript must include the relevant ethics committee approval, consent statement, protocol identifier, and data-use restrictions.

### Funding

[Funding information to be inserted. If no specific funding supported the work, state: The author received no specific funding for this work.]

### Competing interests

[Competing-interest statement to be completed before submission.]

### Author contributions

HZ designed and implemented the software, developed the algorithmic workflow, performed code-level checks, and prepared the manuscript draft. Additional contributors, if any, should be added according to ICMJE and Neuroinformatics authorship criteria before submission.

### Use of generative AI tools

During manuscript preparation, OpenAI Codex was used for drafting assistance, code-audit summarization, and manuscript organization. The author remains responsible for all scientific claims, citations, code behavior, and final submitted content. This statement should be revised to match the final use of AI tools according to Springer Nature and Neuroinformatics policy at the time of submission.

## Information Sharing Statement

SpikeTrainPatternDetector version 1.2.1 is an R/Shiny package with native C components and optional Python integration through reticulate. At submission, the software should be deposited in a public repository with a release tag, license, installation instructions, example data, smoke-test script, dependency versions, and a persistent DOI or URI. The bundled example CSV file inst/extdata/Grechishnikova_STN_2017_subset.csv is included for smoke testing and workflow demonstration; it should not be interpreted as a clinical or disease-physiology dataset. Any private human or animal spike-train recordings analyzed with the software should be described with appropriate institutional approval, consent status, data-access restrictions, and repository information where sharing is permitted."""


def replace_sections(md: str, skip_headings: set[str]) -> str:
    lines = md.splitlines()
    out: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        heading = heading_text(line)
        if heading in skip_headings:
            i += 1
            while i < len(lines):
                next_heading = heading_text(lines[i])
                if lines[i].startswith("## ") and next_heading not in skip_headings:
                    break
                i += 1
            continue
        out.append(line)
        i += 1
    return "\n".join(out).strip() + "\n"


def heading_text(line: str) -> str | None:
    if not line.startswith("## "):
        return None
    text = line.lstrip("#").strip()
    return re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", text)


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_run_font(run, name: str = BASE_FONT, size: float = BASE_SIZE, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def clear_paragraph_borders(style) -> None:
    p_pr = style.element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def set_style_font(style, size: float, bold: bool = False, italic: bool = False, color: str | None = "000000") -> None:
    style.font.name = BASE_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def set_core_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, BASE_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    title = doc.styles["Title"]
    set_style_font(title, 14, bold=True, color="000000")
    clear_paragraph_borders(title)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for name, size, before, after in [
        ("Heading 1", 12, 10, 5),
        ("Heading 2", 11, 7, 4),
        ("Heading 3", 10, 5, 3),
    ]:
        style = doc.styles[name]
        set_style_font(style, size, bold=True, color="000000")
        clear_paragraph_borders(style)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for name in ["Caption", "List Paragraph", "List Bullet", "List Number"]:
        if name in doc.styles:
            style = doc.styles[name]
            set_style_font(style, BASE_SIZE)
            style.paragraph_format.space_after = Pt(3)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_line_numbering(section) -> None:
    sect_pr = section._sectPr
    for existing in list(sect_pr.findall(qn("w:lnNumType"))):
        sect_pr.remove(existing)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "newPage")
    sect_pr.append(ln)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_sep)
    r._r.append(txt)
    r._r.append(fld_end)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        add_line_numbering(section)
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        add_page_number(p)


def add_inline_markdown(paragraph, text: str, base_size: float = BASE_SIZE) -> None:
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=base_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Courier New", size=base_size - 1)
        elif token.startswith("["):
            label_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            run = paragraph.add_run(label_match.group(1) if label_match else token)
            set_run_font(run, size=base_size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size)


def add_para(doc: Document, text: str = "", style: str = "Normal", base_size: float = BASE_SIZE):
    p = doc.add_paragraph(style=style)
    add_inline_markdown(p, text, base_size=base_size)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2:
        header = rows[0]
        body = rows[2:] if all(re.match(r"^-+$|^:?-+:?$", c.replace(" ", "")) for c in rows[1]) else rows[1:]
        return header, body, i
    return [], [], i


def set_cell_text(cell, text: str, bold: bool = False, base_size: float = 8.5) -> None:
    text = text.replace("inst/extdata/Grechishnikova_STN_2017_subset.csv", "Bundled example CSV (inst/extdata)")
    cell.text = ""
    p = cell.paragraphs[0]
    add_inline_markdown(p, text, base_size=base_size)
    for run in p.runs:
        run.bold = bold or run.bold
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_geometry(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_dxa = int(round(sum(widths_in) * 1440))
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    set_table_cell_margins(table)

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width_dxa = int(round(widths_in[min(idx, len(widths_in) - 1)] * 1440))
            cell.width = Inches(widths_in[min(idx, len(widths_in) - 1)])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")


def set_table_cell_margins(table, top=80, bottom=80, start=120, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_word_table(doc: Document, caption: str, header: list[str], body: list[list[str]]) -> None:
    add_para(doc, caption, style="Caption", base_size=BASE_SIZE)
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    ncols = len(header)
    if ncols == 2:
        widths = [3.8, 2.7]
        cell_size = 9
    elif ncols == 4:
        widths = [1.3, 1.85, 1.5, 1.85]
        cell_size = 8.2
    elif ncols == 5:
        widths = [1.0, 1.75, 0.85, 1.25, 1.65]
        cell_size = 7.8
    else:
        widths = [6.5 / max(ncols, 1)] * ncols
        cell_size = 8
    set_table_geometry(table, widths)
    for j, cell in enumerate(table.rows[0].cells):
        set_cell_text(cell, header[j], bold=True, base_size=cell_size)
    for row in body:
        cells = table.add_row().cells
        for j, cell in enumerate(cells):
            set_cell_text(cell, row[j] if j < len(row) else "", base_size=cell_size)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def clean_heading(text: str) -> str:
    heading = text.strip().lstrip("#").strip()
    return re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", heading)


def extract_first_matching(lines: list[str], label: str) -> str:
    for line in lines:
        if line.startswith(label):
            return line.split(":", 1)[1].strip()
    return ""


def clean_metadata_value(value: str) -> str:
    value = value.strip()
    if value.startswith("**"):
        value = value[2:].lstrip()
    if value.endswith("**"):
        value = value[:-2].rstrip()
    return value.strip()


def count_words(md: str) -> int:
    text = re.sub(r"`[^`]+`", " ", md)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)
    return len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", text))


def build() -> None:
    md = transform_markdown(BASE_MD_PATH.read_text(encoding="utf-8"))
    OUT_MD_PATH.write_text(md, encoding="utf-8")
    lines = md.splitlines()

    title = clean_heading(next(line for line in lines if line.startswith("# ")))
    article_type = clean_metadata_value(extract_first_matching(lines, "**Article type:**"))
    target_journal = clean_metadata_value(extract_first_matching(lines, "**Target journal:**"))
    running_title = clean_metadata_value(extract_first_matching(lines, "**Running title:**"))
    authors = clean_metadata_value(extract_first_matching(lines, "**Authors:**"))
    keywords = clean_metadata_value(extract_first_matching(lines, "**Keywords:**"))
    figure_count = sum(1 for line in lines if line.startswith("**Fig. "))
    table_count = sum(1 for line in lines if line.startswith("**Table "))
    word_count = count_words(md)

    doc = Document()
    set_core_styles(doc)
    configure_sections(doc)

    add_para(doc, title, style="Title", base_size=14)
    add_para(doc, authors)
    add_para(doc, "1. [Institution, Department, City, Country - to be completed before submission]")
    add_para(doc, "* Correspondence: Houchun Zhou, zhouhouchun@outlook.com")
    add_para(doc, f"Article type: {article_type}")
    add_para(doc, f"Target journal: {target_journal}")
    add_para(doc, f"Running title: {running_title}")
    add_para(doc, f"Word count: {word_count}; Figures: {figure_count}; Tables: {table_count}")
    add_para(doc, f"Keywords: {keywords}")

    abstract_start = lines.index("## Abstract") + 1
    intro_start = next(i for i, line in enumerate(lines) if line.startswith("## 1. "))
    add_para(doc, "Abstract", style="Heading 1")
    for line in lines[abstract_start:intro_start]:
        if line.strip():
            add_para(doc, line.strip())

    tables_to_append: list[tuple[str, list[str], list[list[str]]]] = []
    i = intro_start
    in_references = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("**Table "):
            caption = re.sub(r"^\*\*|\*\*$", "", stripped)
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and lines[j].strip().startswith("|"):
                header, body, next_i = parse_markdown_table(lines, j)
                tables_to_append.append((caption, header, body))
                add_para(doc, f"{caption} is provided in the Tables section at the end of the manuscript.")
                i = next_i
                continue

        if stripped.startswith("## "):
            heading = clean_heading(stripped)
            in_references = heading == "References"
            add_para(doc, heading, style="Heading 1")
        elif stripped.startswith("### "):
            add_para(doc, clean_heading(stripped), style="Heading 2")
        elif stripped.startswith("#### "):
            add_para(doc, clean_heading(stripped), style="Heading 3")
        elif re.match(r"^\d+\.\s+", stripped):
            p = add_para(doc, re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.space_after = Pt(3)
        elif stripped.startswith("- "):
            p = add_para(doc, stripped[2:], style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.space_after = Pt(3)
        elif stripped.startswith("|"):
            header, body, next_i = parse_markdown_table(lines, i)
            add_word_table(doc, "Table", header, body)
            i = next_i
            continue
        else:
            p = add_para(doc, stripped)
            if in_references and re.match(r"^[A-Z].*\([0-9]{4}\)\.", stripped):
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.paragraph_format.space_after = Pt(4)
        i += 1

    if tables_to_append:
        doc.add_page_break()
        add_para(doc, "Tables", style="Heading 1")
        for caption, header, body in tables_to_append:
            add_word_table(doc, caption, header, body)

    doc.core_properties.title = title
    doc.core_properties.author = "Houchun Zhou"
    doc.core_properties.subject = "Neuroinformatics Software Original Article manuscript"
    doc.core_properties.keywords = keywords
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
