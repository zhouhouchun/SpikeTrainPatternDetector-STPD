from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
TEX_PATH = ROOT / "neuroinformatics_software_original_article_top_draft.tex"
BIB_PATH = ROOT / "references_neuroinformatics.bib"
OUT_PATH = ROOT / "neuroinformatics_software_original_article_top_draft.docx"
AUDIT_PATH = ROOT / "neuroinformatics_software_original_article_top_draft_docx_audit.txt"

BASE_FONT = "Times New Roman"
BASE_SIZE = 10
CONTENT_WIDTH_IN = 6.5


def set_run_font(run, name: str = BASE_FONT, size: float = BASE_SIZE, bold=None, italic=None) -> None:
    run.font.name = name
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float, bold: bool = False, color: str = "000000") -> None:
    style.font.name = BASE_FONT
    if style._element.rPr is not None:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def clear_paragraph_borders(style) -> None:
    p_pr = style.element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, BASE_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    title = doc.styles["Title"]
    set_style_font(title, 14, bold=True)
    clear_paragraph_borders(title)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for name, size, before, after in [
        ("Heading 1", 12, 10, 5),
        ("Heading 2", 11, 7, 4),
        ("Heading 3", 10, 5, 3),
    ]:
        style = doc.styles[name]
        set_style_font(style, size, bold=True)
        clear_paragraph_borders(style)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for name in ["Caption", "List Paragraph", "List Bullet", "List Number"]:
        if name in doc.styles:
            style = doc.styles[name]
            set_style_font(style, BASE_SIZE)
            style.paragraph_format.space_after = Pt(3)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_sep)
    r._r.append(txt)
    r._r.append(fld_end)


def add_line_numbering(section) -> None:
    sect_pr = section._sectPr
    for existing in list(sect_pr.findall(qn("w:lnNumType"))):
        sect_pr.remove(existing)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "newPage")
    sect_pr.append(ln)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        add_line_numbering(section)
        header = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        header.clear()
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = header.add_run("SpikeTrainPatternDetector Neuroinformatics Software Original Article")
        set_run_font(run, size=9)
        footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        footer.clear()
        add_page_number(footer)


def strip_comments(tex: str) -> str:
    out = []
    for line in tex.splitlines():
        if line.lstrip().startswith("%"):
            continue
        out.append(line)
    return "\n".join(out)


def extract_braced_after(tex: str, command: str) -> str:
    pos = tex.find(command)
    if pos < 0:
        return ""
    brace = tex.find("{", pos)
    if brace < 0:
        return ""
    value, _ = read_balanced(tex, brace)
    return value


def read_balanced(text: str, open_pos: int) -> tuple[str, int]:
    assert text[open_pos] == "{"
    depth = 0
    buf = []
    i = open_pos
    while i < len(text):
        ch = text[i]
        if ch == "{" and (i == 0 or text[i - 1] != "\\"):
            depth += 1
            if depth > 1:
                buf.append(ch)
        elif ch == "}" and (i == 0 or text[i - 1] != "\\"):
            depth -= 1
            if depth == 0:
                return "".join(buf), i + 1
            buf.append(ch)
        else:
            buf.append(ch)
        i += 1
    return "".join(buf), i


def parse_bib(path: Path) -> dict[str, dict[str, str]]:
    raw = path.read_text(encoding="utf-8")
    entries: dict[str, dict[str, str]] = {}
    for match in re.finditer(r"@(\w+)\{([^,]+),(.*?)(?=\n@\w+\{|\Z)", raw, flags=re.S):
        key = match.group(2).strip()
        body = match.group(3)
        fields: dict[str, str] = {"entry_type": match.group(1)}
        for line in body.splitlines():
            line = line.strip().rstrip(",")
            if "=" not in line:
                continue
            name, value = line.split("=", 1)
            value = value.strip()
            if value.startswith("{") and value.endswith("}"):
                value = value[1:-1]
            fields[name.strip().lower()] = value.replace("{{", "{").replace("}}", "}")
        entries[key] = fields
    return entries


def surname(author: str) -> str:
    return author.split(",", 1)[0].strip()


def initials(given: str) -> str:
    parts = re.findall(r"[A-Za-z]+", given)
    return " ".join(f"{p[0]}." for p in parts)


def apa_authors(authors: str) -> str:
    names = [a.strip() for a in authors.split(" and ") if a.strip()]
    converted = []
    for name in names:
        if "," in name:
            last, given = [p.strip() for p in name.split(",", 1)]
            converted.append(f"{last}, {initials(given)}".strip())
        else:
            converted.append(name)
    if len(converted) <= 1:
        return converted[0] if converted else ""
    if len(converted) == 2:
        return f"{converted[0]}, & {converted[1]}"
    return f"{', '.join(converted[:-1])}, & {converted[-1]}"


def cite_text(key: str, bib: dict[str, dict[str, str]]) -> str:
    entry = bib.get(key, {})
    authors = [a.strip() for a in entry.get("author", "").split(" and ") if a.strip()]
    year = entry.get("year", "n.d.")
    if not authors:
        return year
    if len(authors) == 1:
        who = surname(authors[0])
    elif len(authors) == 2:
        who = f"{surname(authors[0])} & {surname(authors[1])}"
    else:
        who = f"{surname(authors[0])} et al."
    return f"{who}, {year}"


def replace_citations(text: str, bib: dict[str, dict[str, str]]) -> str:
    def chen_style(match: re.Match) -> str:
        label = match.group(1)
        key = match.group(2).split(",")[0].strip()
        year = bib.get(key, {}).get("year", "n.d.")
        return f"{label} ({year})"

    text = re.sub(r"([A-Z][A-Za-z-]+ et al\.)\s*\\citep\{([^}]+)\}", chen_style, text)

    def repl(match: re.Match) -> str:
        keys = [k.strip() for k in match.group(1).split(",")]
        return "(" + "; ".join(cite_text(k, bib) for k in keys) + ")"

    return re.sub(r"\\citep\{([^}]+)\}", repl, text)


def replace_frac(text: str) -> str:
    while "\\frac" in text:
        pos = text.find("\\frac")
        first = text.find("{", pos)
        if first < 0:
            break
        num, after_num = read_balanced(text, first)
        second = text.find("{", after_num)
        if second < 0:
            break
        den, after_den = read_balanced(text, second)
        repl = f"({replace_frac(num)})/({replace_frac(den)})"
        text = text[:pos] + repl + text[after_den:]
    return text


def latex_formula_to_text(expr: str) -> str:
    text = expr.strip()
    text = replace_frac(text)
    text = text.replace("\\left", "").replace("\\right", "")
    replacements = {
        "\\Delta": "Δ",
        "\\theta": "θ",
        "\\delta": "δ",
        "\\operatorname{first\\_finite}": "first_finite",
        "\\operatorname{sd}": "sd",
        "\\operatorname{mean}": "mean",
        "\\operatorname{IoU}": "IoU",
        "\\operatorname{Precision}": "Precision",
        "\\operatorname{Recall}": "Recall",
        "\\operatorname{clip}": "clip",
        "\\mathcal{D}": "D",
        "\\mathcal{I}": "I",
        "\\mathrm{ISI}": "ISI",
        "\\mathrm{CV}": "CV",
        "\\mathrm{LV}": "LV",
        "\\mathrm{MM}": "MM",
        "\\mathrm{TP}": "TP",
        "\\mathrm{FP}": "FP",
        "\\mathrm{FN}": "FN",
        "\\bar{T}": "T_bar",
        "\\ast": "*",
        "\\le": "≤",
        "\\ge": "≥",
        "\\ne": "≠",
        "\\in": " ∈ ",
        "\\forall": " ∀ ",
        "\\cap": " ∩ ",
        "\\cup": " ∪ ",
        "\\min": "min",
        "\\max": "max",
        "\\sum": "∑",
        "\\cdots": "...",
        "\\ldots": "...",
        "\\qquad": ", ",
        "\\quad": ", ",
        "\\,": " ",
        "\\;": " ",
        "\\!": "",
        "\\{": "{",
        "\\}": "}",
        "\\_": "_",
        "\\texttt{possible\\_burst}": "possible_burst",
        "\\texttt{": "",
        "\\mathbf{1}": "1",
        "\\emptyset": "∅",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\\mathcal\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\mathrm\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\operatorname\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\mathbf\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\text\{([^}]+)\}", r"\1", text)
    for _ in range(4):
        text = re.sub(r"_\{([^{}]+)\}", r"_\1", text)
        text = re.sub(r"\^\{([^{}]+)\}", r"^\1", text)
    text = text.replace("\\", "")
    text = text.replace("&", "")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r",\s*,", ",", text)
    text = text.replace(" ,", ",").replace("( ", "(").replace(" )", ")")
    return text.strip()


def inline_latex_to_text(text: str, bib: dict[str, dict[str, str]]) -> str:
    text = replace_citations(text, bib)
    text = re.sub(r"\\stpd\{\}", "SpikeTrainPatternDetector", text)
    text = re.sub(r"\\stpd\b", "SpikeTrainPatternDetector", text)
    text = re.sub(r"\\isi\b", "ISI", text)
    text = re.sub(r"\\textsc\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\texttt\{([^}]+)\}", lambda m: m.group(1).replace("\\_", "_"), text)
    text = re.sub(r"\\textbf\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\$(.*?)\$", lambda m: latex_formula_to_text(m.group(1)), text)
    text = text.replace("\\&", "&").replace("\\%", "%").replace("\\_", "_")
    text = re.sub(r"\\[a-zA-Z]+\*?\{([^{}]*)\}", r"\1", text)
    text = text.replace("{}", "")
    return re.sub(r"\s+", " ", text).strip()


def add_para(doc: Document, text: str = "", style: str = "Normal", size: float = BASE_SIZE):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_highlight(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=BASE_SIZE, bold=True)
    run.font.color.rgb = RGBColor.from_string("C00000")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.paragraph_format.space_after = Pt(6)


def add_equation(doc: Document, expr: str, number: int) -> None:
    label = doc.add_paragraph(style="Caption")
    r = label.add_run(f"Equation ({number})")
    set_run_font(r, size=BASE_SIZE, bold=True)
    label.paragraph_format.keep_with_next = True

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    run = p.add_run(latex_formula_to_text(expr))
    set_run_font(run, name="Cambria Math", size=10.5)


def set_table_cell_margins(table, top=80, bottom=80, start=120, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for margin, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(round(sum(widths_in) * 1440))))
    tbl_w.set(qn("w:type"), "dxa")
    set_table_cell_margins(table)

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_in[min(idx, len(widths_in) - 1)]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(round(width * 1440))))
            tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, caption: str, rows: list[list[str]]) -> None:
    add_para(doc, caption, style="Caption")
    if not rows:
        return
    ncols = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    if ncols == 2:
        widths = [2.9, 3.6]
        cell_size = 8.8
    else:
        widths = [CONTENT_WIDTH_IN / ncols] * ncols
        cell_size = 8.2
    set_table_geometry(table, widths)
    for j, cell in enumerate(table.rows[0].cells):
        set_cell_text(cell, rows[0][j] if j < len(rows[0]) else "", bold=True, size=cell_size)
    for row in rows[1:]:
        cells = table.add_row().cells
        for j, cell in enumerate(cells):
            set_cell_text(cell, row[j] if j < len(row) else "", size=cell_size)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def parse_latex_table(block: str, bib: dict[str, dict[str, str]]) -> tuple[str, list[list[str]]]:
    caption = extract_braced_after(block, "\\caption")
    caption = inline_latex_to_text(caption, bib)
    tab = re.search(r"\\begin\{tabular\}.*?\n(.*?)\\end\{tabular\}", block, flags=re.S)
    rows: list[list[str]] = []
    if not tab:
        return caption, rows
    for line in tab.group(1).splitlines():
        line = line.strip()
        if not line or line.startswith("\\toprule") or line.startswith("\\midrule") or line.startswith("\\bottomrule"):
            continue
        if "&" not in line:
            continue
        line = re.sub(r"\\\\\s*$", "", line)
        cells = [inline_latex_to_text(c.strip(), bib) for c in line.split("&")]
        rows.append(cells)
    return caption, rows


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name="Courier New", size=8.5)


def add_references(doc: Document, bib: dict[str, dict[str, str]]) -> None:
    add_para(doc, "References", style="Heading 1")
    entries = sorted(bib.values(), key=lambda e: (surname(e.get("author", "")), e.get("year", "")))
    for entry in entries:
        authors = apa_authors(entry.get("author", ""))
        year = entry.get("year", "n.d.")
        title = entry.get("title", "").replace("{", "").replace("}", "")
        journal = entry.get("journal", "")
        volume = entry.get("volume", "")
        number = entry.get("number", "")
        pages = entry.get("pages", "").replace("--", "–")
        doi = entry.get("doi", "")
        url = entry.get("url", "")
        vol = volume
        if number:
            vol += f"({number})"
        ref = f"{authors} ({year}). {title}."
        if journal:
            ref += f" {journal}"
        if vol:
            ref += f", {vol}"
        if pages:
            ref += f", {pages}"
        ref += "."
        if doi:
            ref = f"{ref} https://doi.org/{doi}"
        elif url:
            ref = f"{ref} {url}"
        p = add_para(doc, ref)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)


def build_docx() -> None:
    tex = strip_comments(TEX_PATH.read_text(encoding="utf-8"))
    bib = parse_bib(BIB_PATH)

    title = inline_latex_to_text(extract_braced_after(tex, "\\title"), bib)
    title = re.sub(r"^\[[^\]]+\]", "", title).strip()
    abstract = inline_latex_to_text(extract_braced_after(tex, "\\abstract"), bib)
    keywords = inline_latex_to_text(extract_braced_after(tex, "\\keywords"), bib)

    doc = Document()
    configure_styles(doc)
    configure_sections(doc)

    add_para(doc, title, style="Title", size=14)
    add_para(doc, "Houchun Zhou*")
    add_para(doc, "Department to be completed, Institution to be completed, City, Country")
    add_para(doc, "* Correspondence: zhouhouchun@outlook.com")
    add_para(doc, "Article type: Software Original Article")
    add_para(doc, "Target journal: Neuroinformatics")
    add_para(doc, "Running title: Auditable spike-train event review")
    add_para(doc, f"Keywords: {keywords}")

    add_para(doc, "Abstract", style="Heading 1")
    add_para(doc, abstract)

    body = tex.split("\\maketitle", 1)[1].split("\\bibliography", 1)[0]
    lines = body.splitlines()
    paragraph: list[str] = []
    equation_number = 0
    i = 0

    def flush_paragraph() -> None:
        nonlocal paragraph
        if not paragraph:
            return
        text = " ".join(part.strip() for part in paragraph if part.strip())
        text = inline_latex_to_text(text, bib)
        if text:
            add_para(doc, text)
        paragraph = []

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            flush_paragraph()
            i += 1
            continue

        section = re.match(r"\\section\*?\{(.+)\}", line)
        subsection = re.match(r"\\subsection\*?\{(.+)\}", line)
        subsubsection = re.match(r"\\subsubsection\*?\{(.+)\}", line)
        todo = re.match(r"\\(TODOFIG|TODOTAB|TODONOTE)\{(.+)\}", line)

        if line.startswith("\\begin{equation}"):
            flush_paragraph()
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("\\end{equation}"):
                block.append(lines[i])
                i += 1
            equation_number += 1
            add_equation(doc, "\n".join(block), equation_number)
            i += 1
            continue

        if line.startswith("\\begin{lstlisting}"):
            flush_paragraph()
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("\\end{lstlisting}"):
                block.append(lines[i].rstrip())
                i += 1
            add_code_block(doc, "\n".join(block))
            i += 1
            continue

        if line.startswith("\\begin{table}"):
            flush_paragraph()
            block = [line]
            i += 1
            while i < len(lines):
                block.append(lines[i])
                if lines[i].strip().startswith("\\end{table}"):
                    break
                i += 1
            caption, rows = parse_latex_table("\n".join(block), bib)
            add_table(doc, caption, rows)
            i += 1
            continue

        if todo:
            flush_paragraph()
            label = {
                "TODOFIG": "FIGURE PLACEHOLDER",
                "TODOTAB": "TABLE PLACEHOLDER",
                "TODONOTE": "AUTHOR ACTION REQUIRED",
            }[todo.group(1)]
            add_highlight(doc, f"[{label}: {inline_latex_to_text(todo.group(2), bib)}]")
            i += 1
            continue

        if subsubsection:
            flush_paragraph()
            add_para(doc, inline_latex_to_text(subsubsection.group(1), bib), style="Heading 3")
            i += 1
            continue

        if subsection:
            flush_paragraph()
            add_para(doc, inline_latex_to_text(subsection.group(1), bib), style="Heading 2")
            i += 1
            continue

        if section:
            flush_paragraph()
            add_para(doc, inline_latex_to_text(section.group(1), bib), style="Heading 1")
            i += 1
            continue

        if line.startswith("\\end{document}") or line.startswith("\\begin{document}"):
            i += 1
            continue

        paragraph.append(line)
        i += 1

    flush_paragraph()
    add_references(doc, bib)

    doc.core_properties.title = title
    doc.core_properties.author = "Houchun Zhou"
    doc.core_properties.subject = "Neuroinformatics Software Original Article manuscript"
    doc.core_properties.keywords = keywords
    doc.save(OUT_PATH)

    audit_docx(OUT_PATH, equation_number, len(bib))
    print(OUT_PATH)


def audit_docx(path: Path, expected_equations: int, expected_refs: int) -> None:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    equation_count = sum(1 for text in paragraphs if text.startswith("Equation ("))
    heading_count = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))
    highlighted = sum(1 for text in paragraphs if "PLACEHOLDER" in text or "AUTHOR ACTION REQUIRED" in text)
    table_count = len(doc.tables)
    with zipfile.ZipFile(path) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
        title_border = "w:pBdr" in document_xml[:5000]
        raw_latex_commands = re.findall(r"\\[A-Za-z]+", document_xml)
    report = [
        f"docx={path}",
        "preset=standard_business_brief with Neuroinformatics manuscript override",
        "page=US Letter, 1 inch margins, line numbering per page, centered page numbers",
        f"base_font={BASE_FONT}",
        f"base_size_pt={BASE_SIZE}",
        f"equations_expected={expected_equations}",
        f"equations_found={equation_count}",
        f"references_expected={expected_refs}",
        f"headings_found={heading_count}",
        f"tables_found={table_count}",
        f"highlighted_action_items_found={highlighted}",
        f"title_border_present={title_border}",
        f"raw_latex_command_like_tokens={len(raw_latex_commands)}",
    ]
    AUDIT_PATH.write_text("\n".join(report) + "\n", encoding="utf-8")


if __name__ == "__main__":
    build_docx()
