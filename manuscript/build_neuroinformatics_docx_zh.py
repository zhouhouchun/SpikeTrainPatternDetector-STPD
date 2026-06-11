from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))
from build_neuroinformatics_docx_from_tex import parse_bib, apa_authors, surname  # noqa: E402


BIB_PATH = ROOT / "references_neuroinformatics.bib"
OUT_PATH = ROOT / "neuroinformatics_software_original_article_top_draft.zh.docx"
AUDIT_PATH = ROOT / "neuroinformatics_software_original_article_top_draft_zh_docx_audit.txt"

ZH_FONT = "Songti SC"
LATIN_FONT = "Times New Roman"
BASE_SIZE = 10.5
CONTENT_WIDTH_IN = 6.5


def set_run_font(run, size: float = BASE_SIZE, bold=None, italic=None, latin: str = LATIN_FONT) -> None:
    run.font.name = latin
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)
        run._element.rPr.rFonts.set(qn("w:ascii"), latin)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float, bold: bool = False, color: str = "000000") -> None:
    style.font.name = LATIN_FONT
    if style._element.rPr is not None:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def clear_paragraph_borders(style) -> None:
    p_pr = style.element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, BASE_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    title = doc.styles["Title"]
    set_style_font(title, 15, bold=True)
    clear_paragraph_borders(title)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for name, size, before, after in [
        ("Heading 1", 12.5, 10, 5),
        ("Heading 2", 11.5, 7, 4),
        ("Heading 3", 10.5, 5, 3),
    ]:
        style = doc.styles[name]
        set_style_font(style, size, bold=True)
        clear_paragraph_borders(style)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for name in ["Caption", "List Paragraph", "List Bullet", "List Number"]:
        if name in doc.styles:
            style = doc.styles[name]
            set_style_font(style, BASE_SIZE)
            style.paragraph_format.space_after = Pt(3)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_sep)
    r._r.append(txt)
    r._r.append(fld_end)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9)


def add_line_numbering(section) -> None:
    sect_pr = section._sectPr
    for existing in list(sect_pr.findall(qn("w:lnNumType"))):
        sect_pr.remove(existing)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "newPage")
    sect_pr.append(ln)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        add_line_numbering(section)
        header = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        header.clear()
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = header.add_run("SpikeTrainPatternDetector Neuroinformatics Software Original Article 中文稿")
        set_run_font(run, size=9)
        footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        footer.clear()
        add_page_number(footer)


def add_para(doc: Document, text: str = "", style: str = "Normal", size: float = BASE_SIZE):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_highlight(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=BASE_SIZE, bold=True)
    run.font.color.rgb = RGBColor.from_string("C00000")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.paragraph_format.space_after = Pt(6)


def add_equation(doc: Document, number: int, expr: str) -> None:
    label = doc.add_paragraph(style="Caption")
    r = label.add_run(f"公式（{number}）")
    set_run_font(r, size=BASE_SIZE, bold=True)
    label.paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    run = p.add_run(expr)
    set_run_font(run, size=10.5, latin="Cambria Math")


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=8.5, latin="Courier New")


def set_table_cell_margins(table, top=80, bottom=80, start=120, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for margin, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(round(sum(widths_in) * 1440))))
    tbl_w.set(qn("w:type"), "dxa")
    set_table_cell_margins(table)
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_in[min(idx, len(widths_in) - 1)]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(round(width * 1440))))
            tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.7) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, caption: str, rows: list[list[str]]) -> None:
    add_para(doc, caption, style="Caption")
    ncols = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    widths = [2.7, 3.8] if ncols == 2 else [CONTENT_WIDTH_IN / ncols] * ncols
    set_table_geometry(table, widths)
    for j, cell in enumerate(table.rows[0].cells):
        set_cell_text(cell, rows[0][j] if j < len(rows[0]) else "", bold=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for j, cell in enumerate(cells):
            set_cell_text(cell, row[j] if j < len(row) else "")
    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_references(doc: Document, bib: dict[str, dict[str, str]]) -> None:
    add_para(doc, "参考文献", style="Heading 1")
    entries = sorted(bib.values(), key=lambda e: (surname(e.get("author", "")), e.get("year", "")))
    for entry in entries:
        authors = apa_authors(entry.get("author", ""))
        year = entry.get("year", "n.d.")
        title = entry.get("title", "").replace("{", "").replace("}", "")
        journal = entry.get("journal", "")
        volume = entry.get("volume", "")
        number = entry.get("number", "")
        pages = entry.get("pages", "").replace("--", "–")
        doi = entry.get("doi", "")
        url = entry.get("url", "")
        vol = volume + (f"({number})" if number else "")
        ref = f"{authors} ({year}). {title}."
        if journal:
            ref += f" {journal}"
        if vol:
            ref += f", {vol}"
        if pages:
            ref += f", {pages}"
        ref += "."
        if doi:
            ref += f" https://doi.org/{doi}"
        elif url:
            ref += f" {url}"
        p = add_para(doc, ref, size=9.2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)


EQS = [
    "t_1 < t_2 < ... < t_n",
    "Δ_i = t_i − t_{i−1}, i = 2,...,n",
    "c = [a,b], 2 ≤ a ≤ b ≤ n",
    "Δ_i ≥ θ_art",
    "θ_art ≤ Δ_i < θ_ref",
    "θ*_{p,f} = first_finite(θ_user_{p,f}, θ_manual_{p,f}, θ_hist_{p,f}, θ_default_{p,f})",
    "0 ≤ θ*_{p,lower} < θ*_{p,upper} ≤ θ*_{p,bridge}",
    "D(c) = {Δ_a, Δ_{a+1}, ..., Δ_b}",
    "N_ISI(c) = b − a + 1",
    "N_spk(c) = b − a + 2",
    "T(c) = t_b − t_{a−1}",
    "E(c) = [Q_10(c), Q_50(c), Q_90(c), Q_95(c), CV(c), LV(c), MM(c), B(c), F_B(c), S_min(c), R(c)]",
    "CV(c) = sd(D(c))/mean(D(c))",
    "LV(c) = (1/(m−1)) Σ_{j=1}^{m−1} 3(δ_j−δ_{j+1})²/(δ_j+δ_{j+1})²",
    "MM(c) = max(D(c))/mean(D(c))",
    "θ*_{seed,low} ≤ Δ_i ≤ θ*_{seed,high}",
    "θ*_{seed,high} < Δ_i ≤ θ*_{bridge}",
    "F_B(c) = B(c)/N_ISI(c)",
    "S_pre(c) = Δ_{a−1}/Q_90(c), S_post(c) = Δ_{b+1}/Q_90(c)",
    "S_min(c) = min{S_pre(c), S_post(c)}",
    "F_short(c) = (1/N_ISI(c)) Σ_{i=a}^{b} 1{Δ_i ≤ θ_short}",
    "C* = arg max_C Σ_{c∈C} V(c), subject to c_i ∩ c_j = ∅ for all i ≠ j",
    "Eventness(c) = mean[min{C_edge(c), C_ctx(c)}, C_return(c)]",
    "Regularity(c) = (1/3)[1 − clip_[0,1](CV(c)/r_CV) + 1 − clip_[0,1](LV(c)/r_LV) + 1 − clip_[0,1]((Q_90(c)/Q_10(c)−1)/(r_Q−1))]",
    "ML = (1/|I_<|) Σ_{i∈I_<} T_i, where I_< = {i: T_i < T_bar}",
    "z_i = log_10(1000T_i)",
    "|A ∩ M| = max(0, min(a_2,m_2) − max(a_1,m_1) + 1)",
    "|A ∪ M| = (a_2−a_1+1) + (m_2−m_1+1) − |A ∩ M|",
    "IoU(A,M) = |A ∩ M|/|A ∪ M|",
    "Precision = TP/(TP+FP), Recall = TP/(TP+FN)",
    "F_1 = 2 Precision Recall/(Precision+Recall)",
    "CV2 = (1/(m−1)) Σ_{j=1}^{m−1} 2|δ_j−δ_{j+1}|/(δ_j+δ_{j+1})",
]


def para(doc: Document, text: str) -> None:
    add_para(doc, text)


def heading(doc: Document, text: str, level: int = 1) -> None:
    add_para(doc, text, style=f"Heading {level}")


def build() -> None:
    bib = parse_bib(BIB_PATH)
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)

    title = "SpikeTrainPatternDetector：用于脉冲序列事件审阅与验证的可审计 R/Shiny 软件框架"
    add_para(doc, title, style="Title", size=15)
    para(doc, "周厚淳*")
    para(doc, "待补充科室，待补充机构，城市，国家")
    para(doc, "* 通讯作者：zhouhouchun@outlook.com")
    para(doc, "文章类型：Software Original Article")
    para(doc, "目标期刊：Neuroinformatics")
    para(doc, "短标题：可审计的脉冲序列事件审阅")
    para(doc, "关键词：脉冲序列；爆发放电检测；事件语法；R/Shiny；可重复神经信息学；人工验证")

    heading(doc, "摘要")
    para(doc, "脉冲序列模式标注将神经元放电动力学与环路状态、行为、病理和刺激效应联系起来，但爆发放电、暂停、张力性放电和高频放电等标签容易受到时间戳伪迹、 spike sorting 质量、操作阈值和人工审阅实践的影响。本文介绍 SpikeTrainPatternDetector，一个面向质量控制和审计追踪的 R/Shiny 脉冲序列事件审阅框架。该软件导入原始或带注释的脉冲时间戳，根据时间戳重新计算相邻脉冲间隔，阻止将派生输出误作为原始数据重复分析，在检测前执行质量控制，并根据用户设置、人工标签、直方图结构或默认值解析检测阈值，同时导出带参数哈希的运行元数据。其核心检测器是一套具有生理动机的事件语法，可基于短间隔种子、桥接间隔、侧翼对比、局部规则性、暂停上下文和加权区间选择生成可审阅的候选事件。Mean-ISI 和对数 ISI 直方图支持模块作为独立阈值证据层提供，而不是隐藏的最终标签来源。每个自动事件均可追溯至候选账本、诊断审计、事件性特征、阈值表、最终决策记录和人工标签验证指标。SpikeTrainPatternDetector 旨在作为透明的候选生成与验证平台，而不是通用的自动真值分类器。")

    heading(doc, "引言")
    para(doc, "动作电位是离散事件，但许多神经科学问题关注的是放电模式片段或状态。研究者使用爆发放电、暂停、张力性放电、高频脉冲包以及这些状态之间的转换来描述神经元活动，因为这些描述能够把脉冲时序与突触驱动、环路状态、病理动力学、刺激效应和行为联系起来。在基底节生理、运动系统、皮层网络记录和细胞外单单位研究中，这些标签具有明确的科学价值；但它们也具有方法学脆弱性。同一条脉冲序列可能因最小有效间隔、重复时间戳、疑似不应期间隔、非平稳放电和候选审阅策略不同而得到不同标签。")
    para(doc, "已有方法解决了这一问题的不同侧面。变异系数、局部变异和相邻间隔变异系数 CV2 可概括放电不规则性和局部时间结构；Mean-ISI 与 logISIH/newBD 可提供有文献依据的自适应阈值证据；PCA、因子分析、GPFA、Isomap、扩散映射、t-SNE、UMAP、PHATE、CEBRA 和 sliceTCA 等方法则支持群体数据的低维结构探索。SpikeTrainPatternDetector 的目标不是替代这些方法，而是把日常脉冲序列事件审阅中最容易出错的环节变得透明、可追溯和可验证。")
    para(doc, "实际分析中，研究者需要的不仅是一个分类器，还需要能够导入原始时间戳、避免派生表误导入、在检测前执行质量控制、生成可解释候选标签、保护人工标注、暴露失败模式、记录参数选择、比较自动结果与人工标签并导出可重复证据的软件。SpikeTrainPatternDetector 正是在这一语境下开发的：它以事件语法提出可审阅候选，以文献支持模块提供阈值证据，以导出表格和验证指标支持审稿人与用户追踪每个标签的来源。")
    add_highlight(doc, "[图 1 占位：建议在引言末尾放置总体流程图，从原始时间戳、质量控制、阈值解析、事件语法、候选账本、验证报告到可选探索模块。]")
    add_highlight(doc, "[作者动作：图件需符合 Neuroinformatics/Springer 要求：文件名按 Fig1.eps 等连续命名；矢量图优先 EPS，半色调图用 TIFF；线图 1200 dpi、组合图 600 dpi、半色调图 300 dpi；RGB 8 bits/channel；Helvetica 或 Arial 8–12 pt；图内不放标题；除颜色外使用图案，并保证文字对比度不低于 4.5:1。]")

    heading(doc, "方法")
    heading(doc, "软件实现和设计原则", 2)
    para(doc, "SpikeTrainPatternDetector 主要用 R 实现，提供 Shiny 图形界面和 Plotly 交互式可视化。原生 C 后端加速重复的低层操作，包括单序列 ISI 分位数、局部中位数缓存、短 ISI 连续段扫描、结构候选扫描和区间重叠计算。软件暴露版本中立的公共入口，包括构建数据集、运行检测器、导出结果和计算事件级验证摘要等函数。检测参数由 YAML 支撑的参数契约物化，该契约定义默认值、公共与内部参数命名空间、验证约束和用户界面元数据。")
    para(doc, "设计遵循五个原则：第一，在计算任何模式证据前检查原始脉冲时序完整性；第二，候选生成与公开事件调用分离，诊断窗口、被拒候选和 profile-only 摘要保留用于审计，但不计入生物学事件；第三，阈值在数据集入口一次性解析并在运行中冻结，使事件调用可由阈值表和参数哈希复现；第四，核心检测器被描述为操作性事件语法，而不是普适真值分类器；第五，人工标签、基准测试和参数敏感性分析是性能主张的必要证据。")
    heading(doc, "安装和最小运行示例", 2)
    para(doc, "软件应作为标准 R 包从带版本的源代码仓库或归档版本安装。最小程序化流程如下：")
    add_code(doc, 'library(SpikeTrainPatternDetector)\n\nparams <- default_params()\nds <- build_spike_dataset("spikes.csv", mode = "raw", unit_in = "s")\nds <- stpd_detect(ds, params,\n                  lock_manual = TRUE,\n                  collect_diagnostics = TRUE)\nstpd_export_results(ds, params, out_dir = "results")')
    para(doc, "图形界面提供相同的导入、质量控制、阈值审阅、候选检查、人工标签验证和导出流程。正式分析应报告包版本、源代码提交或 release tag、R 版本、操作系统、参数哈希、阈值表、质量控制表，以及任何可选 Python 后端版本。")

    heading(doc, "输入数据模型", 2)
    para(doc, "主要输入是一张脉冲时间戳表。在原始 CSV 模式下，每个非空数值列被解释为一条脉冲序列。时间戳可以秒或毫秒给出，并在内部转换为秒。导入器移除缺失值、对每条序列内时间戳排序、记录输入顺序是否非单调，并由时间戳差值重新计算所有相邻脉冲间隔。对有序时间戳：")
    add_equation(doc, 1, EQS[0])
    para(doc, "检测器定义：")
    add_equation(doc, 2, EQS[1])
    para(doc, "所有候选事件都表示为连续的 ISI 索引区间：")
    add_equation(doc, 3, EQS[2])
    para(doc, "这种表示使事件边界显式化，并允许自动事件与人工事件通过区间重叠比较。在带注释输入模式中，时间戳列和人工标签列可以同时导入。外部 ISI 列若存在，只用于质量控制检查，不作为检测依据。候选账本、事件性审计、最终事件表、阈值表或诊断候选表等派生输出在原始导入中默认被阻止，以降低把检测器输出递归当作原始数据的风险。")

    heading(doc, "检测前质量控制", 2)
    para(doc, "质量控制在爆发、暂停、张力性或高频证据计算前完成。每条序列的 QC 表报告脉冲数、记录时长、放电率、原始最小 ISI、最小有效 ISI、精确重复时间戳、零或负时间步、零或负 ISI、硬伪迹 ISI、疑似不应期 ISI、时间戳与 ISI 是否不一致以及平稳性诊断。默认硬伪迹阈值为 0.0009 s；第二个疑似不应期阈值默认为 0.001 s。有效正 ISI 满足：")
    add_equation(doc, 4, EQS[3])
    para(doc, "疑似不应期间隔满足：")
    add_equation(doc, 5, EQS[4])
    para(doc, "默认情况下，重复时间戳、零或负 ISI 以及硬伪迹 ISI 会在生成任何模式标签前停止检测。探索性运行可以显式允许在 QC 错误后继续，但正式分析应在导入时合并精确重复时间戳，或报告其比例和敏感性影响。平稳性警告作为解释风险导出，而不是硬错误。")

    heading(doc, "参数物化与阈值解析", 2)
    para(doc, "检测器参数存储于 YAML 支撑的参数契约中。公共参数命名空间为 spiketrainpattern，event_core、event_grammar、highfreq、tonic、pause 和 classification 等内部命名空间在检测前派生。每次运行记录参数哈希并导出参数报告。")
    para(doc, "每个模式族 p 和阈值字段 f 的候选值可以来自用户设置、人工标签摘要、直方图建议或默认值。活动阈值可写为：")
    add_equation(doc, 6, EQS[5])
    para(doc, "随后施加几何约束：")
    add_equation(doc, 7, EQS[6])
    para(doc, "导出的阈值表记录每个模式和字段的所有候选值、最终有效值和来源，是复现运行的关键证据。")
    add_highlight(doc, "[表 1 占位：软件元数据、实现语言、许可证、操作系统说明、依赖、公共 API、示例数据和仓库/DOI 字段。]")

    heading(doc, "事件语法检测器", 2)
    para(doc, "核心检测器是在 ISI 索引候选区间上的确定性事件语法。对候选区间 c=[a,b]，定义候选内 ISI 集合、ISI 数、脉冲数和时长如下：")
    for n in range(8, 12):
        add_equation(doc, n, EQS[n - 1])
    para(doc, "检测器计算证据向量：")
    add_equation(doc, 12, EQS[11])
    para(doc, "其中 Q_q(c) 为候选内第 q 分位数，CV 为变异系数，LV 为局部变异，MM 为最大值与均值比，B(c) 为桥接数，F_B(c) 为桥接比例，S_min(c) 为最小侧翼对比，R(c) 概括规则性或状态族证据。对有效正 ISI，有：")
    for n in range(13, 16):
        add_equation(doc, n, EQS[n - 1])

    heading(doc, "爆发族候选", 3)
    para(doc, "爆发族检测从紧凑短 ISI 种子段开始。若某个 ISI 满足：")
    add_equation(doc, 16, EQS[15])
    para(doc, "则它支持种子。连续种子间隔形成种子段，并需达到配置的最小种子 ISI 数。种子段可向左右通过有限数量的桥接间隔扩展，桥接间隔满足：")
    add_equation(doc, 17, EQS[16])
    para(doc, "扩展受最大桥接数、桥接比例和最大扩展步数约束。桥接比例为：")
    add_equation(doc, 18, EQS[17])
    para(doc, "当两侧侧翼均可观测时，候选爆发族区间按侧翼 ISI 评估：")
    add_equation(doc, 19, EQS[18])
    add_equation(doc, 20, EQS[19])
    para(doc, "典型双侧爆发族通过要求两个侧翼比值均超过解析得到的对比阈值，并要求桥接负担和事件内紧凑性保持在配置限制内。单侧或边界受限候选可以保留为 possible_burst，而不是直接丢弃。")

    heading(doc, "高频放电、张力性状态和暂停", 3)
    para(doc, "高频放电被建模为持续状态或时段，而非爆发族事件。候选高频放电时段由多数 ISI 较短、同时容忍少量中等间隔的支持段构成，其短 ISI 比例为：")
    add_equation(doc, 21, EQS[20])
    para(doc, "高频张力性和张力性候选表示相对规则的放电状态，分别使用高频上界、中等 ISI 带和 CV、LV、最大均值比等规则性约束。暂停候选表示长间隔，其有效暂停下限结合绝对阈值、序列内上分布和张力性状态防护，以减少将普通低频放电误标为暂停的风险。")
    add_highlight(doc, "[表 2 占位：默认事件语法参数、单位、生物学理由、公式类别和导出审计字段。]")
    add_highlight(doc, "[图 2 占位：展示爆发种子、桥接间隔、侧翼对比、possible_burst 边界情形、高频放电时段和暂停间隔的 ISI 示意图。]")

    heading(doc, "加权区间选择与人工标签策略", 2)
    para(doc, "不同检测层可能产生相互重叠的候选区间。为构建单一公开自动标签轨道，检测器对候选进行加权区间选择。每个候选 c 根据标签族、显式优先级、得分和跨度获得价值 V(c)，最终选择集合为：")
    add_equation(doc, 22, EQS[21])
    para(doc, "这种动态规划式区间选择避免任意的先到先得规则，使爆发、possible_burst、高频、张力性和暂停候选之间的竞争具有确定性。人工标签与自动证据分离处理。启用人工锁定时，人工标注区间保持最终标签优势，自动标签不会覆盖这些 ISI；自动证据仍可在诊断层生成并审计。显式 not-burst 等负人工标签可以否决候选，并保留在诊断审计中。")

    heading(doc, "候选账本与事件性审计", 2)
    para(doc, "被接受并选择的区间导出到公开事件表和候选账本；被拒绝、阻断、profile-only、被抑制或未选中的窗口保留在诊断候选审计中。这种分离防止探索性窗口膨胀公开生物学计数，同时保留失败模式证据。事件性审计通过边缘对比、上下文对比、回归基线证据和规则性判断候选更像离散事件还是状态。事件性得分表示为：")
    add_equation(doc, 23, EQS[22])
    para(doc, "规则性由标准化 CV、LV 和 Q_90/Q_10 证据计算：")
    add_equation(doc, 24, EQS[23])
    para(doc, "事件性是审计特征，不是独立的生物学真值度量。")

    heading(doc, "文献支持方法", 2)
    para(doc, "核心事件语法为项目定义，但软件包含文献支持的阈值证据层。Mean-ISI 支持层遵循 Chen 等（2009）。给定有效 ISI T_i，全局平均 ISI 为 T_bar，Mean-ISI 阈值为：")
    add_equation(doc, 25, EQS[24])
    para(doc, "当窗口内连续 ISI 的均值不超过 ML 时，候选爆发窗口被识别。阈值原则来自文献，而穷举窗口、预算控制和合并规则属于实现选择。logISIH/newBD 支持层实现 Pasquale 风格的 logISI 阈值证据流程（Pasquale 等，2010）：")
    add_equation(doc, 26, EQS[25])
    para(doc, "其中 T_i 以秒计，1000T_i 以毫秒计。软件构建、平滑并搜索 logISI 直方图的峰和谷，使用 void-parameter 证据支持阈值解析。无法解析的阈值会被显式报告为 unresolved，而不会转化为最终标签。")

    heading(doc, "基于人工标签的事件级验证", 2)
    para(doc, "当存在人工标签时，自动事件通过 ISI 索引的交并比与人工事件比较。对自动区间 A=[a_1,a_2] 和人工区间 M=[m_1,m_2]：")
    for n in range(27, 30):
        add_equation(doc, n, EQS[n - 1])
    para(doc, "候选/人工配对按 IoU 从高到低贪婪匹配，并需满足最小重叠标准。精确率、召回率和 F1 为：")
    add_equation(doc, 30, EQS[29])
    add_equation(doc, 31, EQS[30])
    para(doc, "严格高置信度评估保留具体标签；候选族评估可将 burst 和 possible_burst 等相关标签归并，以衡量事件族检索。参数敏感性扫描扰动 Basic 层参数并重算事件级一致性，从而记录精确率、召回率、F1、IoU、边界误差和混淆结构对阈值选择的依赖。")
    add_highlight(doc, "[图 3 占位：人工区间、自动区间、IoU 匹配、严格模式和候选族模式的验证示意图。]")

    heading(doc, "状态空间和群体探索模块", 2)
    para(doc, "SpikeTrainPatternDetector 包含可选的 ISI 状态空间、神经流形、事件对齐活动和 slice tensor 构建模块。这些模块不是核心检测器所必需，也不应单独作为事件为生物学真值的证据。单序列状态空间特征包括 logISI、局部中位 ISI、局部均值 ISI、CV、LV、CV2、滞后 ISI 和局部上下文摘要。CV2 定义为：")
    add_equation(doc, 32, EQS[31])
    para(doc, "线性和非线性嵌入（包括 PCA、Isomap、扩散映射式嵌入、PHATE、UMAP 和 t-SNE）作为可视化工具提供，应报告设置、随机种子和稳定性控制。群体层面可将选定脉冲序列分箱为时间×神经元矩阵，事件标签作为注释叠加，而不作为定义嵌入的隐藏输入，以避免循环推断。")
    add_highlight(doc, "[图 4 占位：如包含群体数据，建议展示群体流形/事件几何工作流，并明确标注其探索性性质，除非已有保持样本解码或行为验证。]")

    heading(doc, "结果")
    heading(doc, "软件工作流和导出审计工件", 2)
    para(doc, "SpikeTrainPatternDetector 实现从原始时间戳到可审阅事件标签和验证报告的可重复工作流。典型运行包括原始或注释导入、QC、重复时间戳处理、阈值解析、事件语法候选生成、候选特征提取、最终决策审计、人工标签可用时的事件级验证，以及公开和诊断表格导出。")
    para(doc, "稳定的公开结果字段包括有效阈值表、检测前质量表、候选诊断审计、候选账本、事件表、候选特征表、最终决策审计、事件性审计、参数报告、参数验证报告、运行元数据和人工标签验证摘要。该组织方式使审阅者能够从最终事件标签回溯到候选窗口、阈值策略、特征证据和参数哈希。")
    add_highlight(doc, "[表 3 占位：列出所有导出文件及其回答的科学问题，包括 Events_final.csv、Candidate_ledger.csv、Candidate_diagnostic_audit.csv、Eventness_audit.csv、Threshold_table.csv、Detector_run_metadata.csv 和验证导出。]")

    heading(doc, "捆绑数据的冒烟测试演示", 2)
    para(doc, "为验证包级工作流，我们在捆绑原始 CSV 文件 inst/extdata/Grechishnikova_STN_2017_subset.csv 上运行 SpikeTrainPatternDetector 1.2.1。该示例是软件冒烟测试，不应解释为临床或疾病生理结果。数据集包含四条脉冲序列，每条 120 个脉冲。检测前 QC 生成四行序列级 QC。检测器产生 116 个公开自动候选事件：67 个 burst、43 个 pause 和 6 个 possible_burst。公开候选账本含 116 行，诊断候选审计含 185 行，候选特征表含 116 行。")
    add_table(doc, "表 1. 捆绑示例数据集的冒烟测试输出摘要。该表展示端到端软件执行和导出结构，不是性能基准。", [
        ["项目", "数值"],
        ["包版本", "1.2.1"],
        ["输入文件", "inst/extdata/Grechishnikova_STN_2017_subset.csv"],
        ["脉冲序列数量", "4"],
        ["每条序列脉冲数", "120, 120, 120, 120"],
        ["QC 行数", "4"],
        ["公开 AUTO 事件数", "116"],
        ["burst 事件", "67"],
        ["pause 事件", "43"],
        ["possible_burst 事件", "6"],
        ["候选账本行数", "116"],
        ["诊断候选审计行数", "185"],
        ["候选特征行数", "116"],
    ])
    para(doc, "冒烟测试说明了两个软件特征：第一，分析链会分别产生公开事件表和诊断审计表；第二，QC 警告保持可见，并且必须在得出生物学结论前解释。冒烟测试可证明工作流能够运行并导出预期工件，但不能证明检测器在金标准基准上的准确性。")
    heading(doc, "验证导向输出", 2)
    para(doc, "软件导出严格模式和候选族模式两类验证。严格验证适合人工标注区间被视为特定事件类别高置信度真值的场景；候选族验证适合关注检测器是否找回更广义生物学事件族的场景。参数敏感性扫描回应了脉冲序列事件分析中的常见弱点：只报告单一参数集而不展示结论是否依赖狭窄阈值选择。")
    add_highlight(doc, "[表 4 占位：基准验证矩阵。行：合成真值、专家人工子集、经典支持方法比较、伪迹压力测试、参数敏感性。列：数据集、主要指标、预期输出、当前状态。]")
    add_highlight(doc, "[图 5 占位：加入基准数据后展示 precision-recall/F1 曲线、IoU 分布和人工或合成验证的边界误差图。]")

    heading(doc, "讨论")
    para(doc, "SpikeTrainPatternDetector 的贡献是提供一个可审计的神经信息学脉冲序列事件审阅框架。它的强项不是宣称存在普适自动真值，而是把通常脆弱的分析链显性化：原始数据在检测前被检查，阈值被解析并冻结，候选窗口与公开调用分离，人工标签可以受到保护，文献支持方法与项目定义的事件语法保持区分，参数报告和验证工件被导出以供检查。")
    heading(doc, "生物学解释", 2)
    para(doc, "事件语法应被操作性解释。burst 调用表示在所选阈值策略下具有足够局部对比的紧凑短 ISI 结构；pause 调用表示符合暂停和上下文规则的长间隔；tonic 和 high-frequency 调用表示由局部速率和规则性证据定义的状态样时间制度。这些标签是否具有生物学意义，取决于实验制备、spike sorting 质量、细胞类型、脑区、疾病或刺激状态以及行为。")
    heading(doc, "与经典爆发检测方法的关系", 2)
    para(doc, "Mean-ISI 和 logISIH/newBD 重要之处在于提供了文献支持的阈值证据，同时也说明不应把任何单一检测器视为普遍最终答案。Mean-ISI 能适应内在 ISI 统计，但会受到非平稳性和长沉默间隔影响；logISIH/newBD 使用直方图结构和谷值证据，但分箱、平滑和峰谷选择会影响阈值解析。SpikeTrainPatternDetector 将这些方法暴露为支持层，使研究者能够与主事件语法比较，而不是静默混合假设。")
    heading(doc, "与相关软件的比较", 2)
    para(doc, "SpikeTrainPatternDetector 与现有神经信息学软件互补。Neo 提供成熟的 Python 电生理数据对象模型，用于多种文件格式的数据处理；其核心贡献是互操作数据表示，而 SpikeTrainPatternDetector 聚焦于质量控制、事件级审阅、冻结阈值表、候选账本和人工标签验证工件。PySpike 提供高效的 Python 脉冲序列同步和距离分析例程；其目标是比较性脉冲序列指标，而 SpikeTrainPatternDetector 围绕单序列和多序列事件注释、审计轨迹和面向审稿人的导出组织。对本文用例而言，其实际优势是工作流层面的可追责性。")
    heading(doc, "局限性", 2)
    para(doc, "第一，核心事件语法是工程校准的规则系统，具有可解释性和可审计性，但不是来自单篇文献的严格公式。第二，性能主张仍需要人工标签或独立终点支持；缺少专家标注、合成真值或行为关联验证时，输出应视为候选事件。第三，QC 警告会改变解释，重复时间戳、疑似不应期间隔、非平稳性和派生文件误导入均可能影响 burst 和 pause 证据。第四，流形和张量模块应保持探索性，除非有群体记录、行为变量、保持样本解码、shuffle 控制和嵌入稳定性分析支持。")
    heading(doc, "未来工作", 2)
    para(doc, "未来工作包括四个方向：扩展参数元数据，使每个参数都标注方法类别、引用键、公式参考、单位约定和验证状态；加入一等基准验证，包括合成事件注入、专家人工标签、评审者间一致性、经典方法比较、伪迹压力测试和参数敏感性曲线；扩展 golden tests，覆盖边界爆发、不应期 doublet、高频张力状态、pause-near-burst 冲突、长爆发、候选账本分离和导出一致性；将 GPFA、CEBRA 和 sliceTCA 等群体工作流与保持样本行为解码、保持神经元预测和跨会话嵌入稳定性结合。")

    heading(doc, "结论")
    para(doc, "SpikeTrainPatternDetector 提供了一个用于脉冲序列模式分析的可审计 R/Shiny 框架。它结合检测前质量控制、可审阅事件语法候选生成、Mean-ISI 和 logISIH/newBD 文献支持层、人工标签保护、验证导出、状态空间分析、群体探索工具和可选 slice tensor 工作流。该软件最适合被理解为透明的候选生成与验证平台；其科学价值在于使脉冲序列模式分析可检查、可重复和可检验，而不是宣称自动获得普适真值。")

    heading(doc, "声明")
    heading(doc, "伦理批准", 2)
    para(doc, "本文描述软件和捆绑冒烟测试数据集，不报告新的人类或动物实验结果。如最终稿件同时包含患者记录、术中数据、动物实验或未发表实验室数据，必须加入相应伦理委员会批准、知情同意、方案编号和数据使用限制。")
    heading(doc, "利益冲突", 2)
    para(doc, "作者声明不存在利益冲突。[提交前请确认或修改。]")
    heading(doc, "基金资助", 2)
    para(doc, "[请填写基金信息。如无专项资助，可写：作者未获得本研究的专项资助。]")
    heading(doc, "作者贡献", 2)
    para(doc, "HZ 设计并实现软件，开发算法工作流，执行代码级检查并准备稿件草案。[如增加合作者，请提交前修订。]")
    heading(doc, "人工智能工具使用声明", 2)
    para(doc, "大型语言模型用于辅助手稿组织、起草和语言编辑。作者已审阅和修订内容，验证科学主张和引用文献，并对最终稿件承担全部责任。[请按最终实际使用情况和期刊政策保留、修订或删除。]")

    heading(doc, "Information Sharing Statement")
    para(doc, "SpikeTrainPatternDetector 1.2.1 是一个包含原生 C 组件并可通过 reticulate 连接可选 Python 后端的 R/Shiny 包。提交时，源代码、安装说明、参数文件、捆绑示例数据、冒烟测试脚本、依赖版本以及生成稿件图表的脚本将通过带版本的仓库提供给审稿人。发表时，软件及相关可重复材料将通过 release tag 和持久归档标识符公开，例如 Zenodo DOI。冒烟测试使用的捆绑 CSV 文件随包分发，路径为 inst/extdata/Grechishnikova_STN_2017_subset.csv；该文件用于展示软件执行流程，不作为临床或疾病生理数据集呈现。复现所报告软件工作流所需的资源不会仅以“upon request”方式提供。")
    add_highlight(doc, "[作者动作：提交前插入最终 GitHub 仓库 URL、release tag、Zenodo DOI 或等效归档标识符、精确依赖版本以及图表生成脚本路径。]")

    heading(doc, "致谢")
    add_highlight(doc, "[作者动作：如适用，请补充科学反馈、数据贡献者、机构或开源依赖的致谢。]")
    add_references(doc, bib)

    doc.core_properties.title = title
    doc.core_properties.author = "Houchun Zhou"
    doc.core_properties.subject = "Neuroinformatics Software Original Article manuscript, Chinese version"
    doc.core_properties.keywords = "脉冲序列; 爆发放电检测; 事件语法; R/Shiny; 可重复神经信息学; 人工验证"
    doc.save(OUT_PATH)
    audit(OUT_PATH)
    print(OUT_PATH)


def audit(path: Path) -> None:
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    bib = parse_bib(BIB_PATH)
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
        raw_latex = re.findall(r"\\[A-Za-z]+", xml)
    report = [
        f"docx={path}",
        "language=Chinese",
        "target=Neuroinformatics Software Original Article review draft",
        f"equations_found={sum(1 for p in doc.paragraphs if p.text.startswith('公式（'))}",
        f"references_expected={len(bib)}",
        f"reference_links_found={text.count('https://doi.org/') + text.count('https://www.jmlr.org/')}",
        f"tables_found={len(doc.tables)}",
        f"highlighted_action_items_found={sum(1 for p in doc.paragraphs if '占位' in p.text or '作者动作' in p.text)}",
        f"has_information_sharing={'Information Sharing Statement' in text}",
        f"raw_latex_command_like_tokens={len(raw_latex)}",
    ]
    AUDIT_PATH.write_text("\n".join(report) + "\n", encoding="utf-8")


if __name__ == "__main__":
    build()
